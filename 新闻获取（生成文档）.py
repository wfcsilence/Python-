#由于新闻联播官网网页变化，该爬虫已失效
import requests
from bs4 import BeautifulSoup
from datetime import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor,Pt
from docx.oxml.ns import qn
import docx
import pandas as pd
import time

headers = {
    'Accept': 'text/html, */*; q=0.01',
    'Referer': 'http://tv.cctv.com/lm/xwlb/',
    'X-Requested-With': 'XMLHttpRequest',
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/71.0.3578.98 Safari/537.36',
}

def href(date):
    """
    用于获取某天新闻联播各条新闻的链接
    :param date: 日期，形如20190101
    :return: href_list: 返回新闻链接的列表
    """
    href_list = []
    response = requests.get('https://tv.cctv.com/lm/xwlb/day/' + str(date) + '.shtml', headers=headers)
    bs_obj = BeautifulSoup(response.text, 'lxml')
    lis = bs_obj.find_all('li')
    for each in lis:
        href_list.append(each.find('a')['href'])
    return href_list

def newsname(url):
    print(url)
    response = requests.get(url, headers=headers, )
    bs_obj = BeautifulSoup(response.content.decode('utf-8'), 'lxml')
    if 'news.cctv.com' in url:
        name = bs_obj.find('div', {'class': 'cnt_nav'}).contents[3].text
        list_str = list(name)
        del list_str[0:4]
        name = "".join(list_str)
    else:
        name = bs_obj.find('div', {'class': 'cnt_nav'}).contents[3].text
        list_str = list(name)
        del list_str[0:4]
        name = "".join(list_str)
    return name

def newstext(url):
    print(url)
    response = requests.get(url, headers=headers, )
    bs_obj = BeautifulSoup(response.content.decode('utf-8'), 'lxml')
    if 'news.cctv.com' in url:
        text = bs_obj.find('div', {'id': 'content_body'}).text
    else:
        text = bs_obj.find('div', {'class': 'cnt_bd'}).text
    return text


def datelist(beginDate, endDate):
    # beginDate, endDate是形如‘20160601’的字符串或datetime格式
    date_l = [datetime.strftime(x, '%Y%m%d') for x in list(pd.date_range(start=beginDate, end=endDate))]
    return date_l

def save_text(date):
    path=str(date)+"新闻.docx"
    f = docx.Document()
    title=f.add_paragraph(str(date)+"新闻")
    title.paragraph_format.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
    title.style.font.size=Pt(16)
    # open(str(date) + '.txt', 'a', encoding='utf-8')
    y=0
    for each in href(date)[1:]:
        print(newsname(each))
        x1=f.add_heading("",level=2)
        y+=1
        run1=x1.add_run(str(y)+'.'+newsname(each))
        run1.font.name=u'微软雅黑'
        run1._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
        #x1.bold=True
        run1.font.color.rgb = RGBColor(205,38,38)
        x2=f.add_paragraph("")
        x2.add_run(newstext(each))
        x2.style.font.size=Pt(11)
        x2.paragraph_format.first_line_indent = x2.style.font.size*2
        #f.add_paragraph('\n')
    f.styles['Normal'].font.name = u'宋体'
    f.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    f.save(path)

print("请输入年份，比如20211001")
st=input("开始:")
fi=input("结束:")
x=0

for date in datelist(st,fi):
    save_text(date)
    x+=1
    print("第"+str(x)+"个文件输出完成")
    time.sleep(2)

print("全部完成，再见~")
time.sleep(2)
